"""
generate_report.py — Generates PROJECT_REPORT.docx from PROJECT_REPORT.md
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy, re

# ── Helpers ────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def add_horizontal_rule(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "B87333")
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_after = Pt(4)


def add_paragraph_style(doc, text, style="Normal", bold=False, italic=False,
                         size=11, color=None, space_before=0, space_after=6, alignment=None):
    p = doc.add_paragraph(style=style)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*bytes.fromhex(color))
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    if alignment:
        p.alignment = alignment
    return p


def add_inline_text(para, text, bold=False, italic=False, size=11, color=None, code=False):
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if code:
        run.font.name = "Courier New"
        run.font.size = Pt(9.5)
    if color:
        run.font.color.rgb = RGBColor(*bytes.fromhex(color))
    return run


def make_table(doc, headers, rows, header_bg="B87333", alt_row_bg="FFF5EC"):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"

    # Header row
    hdr_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(10)
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        cell.paragraphs[0].paragraph_format.space_after = Pt(2)
        set_cell_bg(cell, header_bg)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Data rows
    for ri, row_data in enumerate(rows):
        row = table.rows[ri + 1]
        bg = alt_row_bg if ri % 2 == 0 else "FFFFFF"
        for ci, cell_text in enumerate(row_data):
            cell = row.cells[ci]
            cell.text = str(cell_text)
            cell.paragraphs[0].runs[0].font.size = Pt(9.5)
            cell.paragraphs[0].paragraph_format.space_after = Pt(2)
            set_cell_bg(cell, bg)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    doc.add_paragraph()
    return table


def add_code_block(doc, code_text):
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Inches(0.3)
    para.paragraph_format.space_before = Pt(4)
    para.paragraph_format.space_after = Pt(4)
    # shading
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F5F0EB")
    pPr.append(shd)
    run = para.add_run(code_text)
    run.font.name = "Courier New"
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(0x3A, 0x28, 0x10)
    return para


# ── Document construction ──────────────────────────────────────────────────

doc = Document()

# Page margins
sections = doc.sections
for section in sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ── Title Page ─────────────────────────────────────────────────────────────
doc.add_paragraph()
doc.add_paragraph()

t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("AI-POWERED TRAVEL ITINERARY PLANNER")
r.bold = True
r.font.size = Pt(22)
r.font.color.rgb = RGBColor(0xB8, 0x73, 0x33)

t2 = doc.add_paragraph()
t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = t2.add_run("Mini Project Report — Generative AI")
r2.bold = True
r2.font.size = Pt(14)
r2.font.color.rgb = RGBColor(0x55, 0x44, 0x33)

doc.add_paragraph()
add_horizontal_rule(doc)
doc.add_paragraph()

meta = [
    ("Project Title",   "AI-Powered Travel Itinerary Planner"),
    ("Domain",          "Generative AI, RAG, Full-Stack Web Development"),
    ("Submitted By",    "[Your Name]"),
    ("Institution",     "[Your Institution / College Name]"),
    ("Guide",           "[Guide Name]"),
    ("Date",            "April 2026"),
]
for label, value in meta:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_inline_text(p, f"{label}: ", bold=True, size=12, color="B87333")
    add_inline_text(p, value, size=12)
    p.paragraph_format.space_after = Pt(4)

doc.add_page_break()

# ── Helper to add section heading ─────────────────────────────────────────
def section_heading(num, title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(f"{num}. {title.upper()}")
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(0xB8, 0x73, 0x33)
    add_horizontal_rule(doc)


def sub_heading(title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x55, 0x33, 0x11)


def body(text, indent=False, size=10.5):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.size = Pt(size)
    return p


def bullet(text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.3 + level * 0.2)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    run.font.size = Pt(10.5)


def numbered_item(num, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Inches(0.3)
    p.paragraph_format.space_after  = Pt(3)
    add_inline_text(p, f"{num}. ", bold=True, size=10.5, color="B87333")
    add_inline_text(p, text, size=10.5)


# ── 1. Abstract ──────────────────────────────────────────────────────────
section_heading(1, "Abstract")
body(
    "This project presents an intelligent travel planning web application that generates detailed, "
    "personalized day-by-day travel itineraries using a combination of Retrieval-Augmented Generation (RAG) "
    "and a Large Language Model (LLM) hosted on Groq. The system retrieves structured destination data from "
    "a local ChromaDB vector database, validates user inputs, computes realistic flight costs, allocates "
    "budgets intelligently across trip components, and generates specific, actionable itineraries. "
    "When RAG data is insufficient, the system seamlessly falls back to the LLM's world knowledge. "
    "A React-based frontend provides an intuitive interface for planning and conversational refinement."
)

# ── 2. Problem Statement ──────────────────────────────────────────────────
section_heading(2, "Problem Statement")
body(
    "Existing travel planning tools either produce generic, template-based itineraries or require expensive "
    "real-time API subscriptions. Travelers commonly face:"
)
bullet("Generic recommendations that do not consider personal budget constraints")
bullet("Inability to split a travel budget realistically across flights, stays, food, and activities")
bullet("No differentiation between domestic and international travel costs")
bullet("Failure to validate whether the destination or source city is serviceable")
body(
    "This project addresses these gaps by combining structured travel data from curated CSV datasets "
    "with a state-of-the-art LLM, ensuring both data-grounded accuracy and personalized specificity."
)

# ── 3. Objectives ────────────────────────────────────────────────────────
section_heading(3, "Objectives")
objectives = [
    "Build a full-stack AI application with a FastAPI backend and React frontend",
    "Implement RAG using ChromaDB to ground LLM responses in real destination data",
    "Validate user inputs — source city, destination, and budget — before generation",
    "Compute realistic flight cost estimates with domestic / neighbour / international differentiation",
    "Allocate the post-flight 'ground budget' intelligently across accommodation, food, activities, and transport",
    "Produce specific named recommendations (real hotels, restaurants, landmarks) — not generic placeholders",
    "Handle LLM failures gracefully with a 3-tier JSON repair and fallback mechanism",
]
for i, obj in enumerate(objectives, 1):
    numbered_item(i, obj)

# ── 4. System Architecture ────────────────────────────────────────────────
section_heading(4, "System Architecture")
body(
    "The application follows a layered architecture with clear separation between the frontend UI, "
    "the FastAPI backend processing pipeline, and external AI/data services."
)
add_code_block(doc,
"┌─────────────────────────────────────────────────────┐\n"
"│               React Frontend (Vite)                  │\n"
"│  Travel Form → Itinerary Display → Chat Interface    │\n"
"└──────────────────────┬──────────────────────────────┘\n"
"                       │ HTTP POST /generate  /chat\n"
"┌──────────────────────▼──────────────────────────────┐\n"
"│            FastAPI Backend (test2.py)                │\n"
"│                                                      │\n"
"│  ① Input Validation (Source / Destination in DB?)   │\n"
"│  ② Flight Cost Lookup (RapidAPI → Simulator)         │\n"
"│  ③ Budget Validation (post-flight ground budget)     │\n"
"│  ④ Budget Allocation (Stay/Food/Activities/Transp.)  │\n"
"│  ⑤ RAG Retrieval (ChromaDB vector search, k=6)      │\n"
"│  ⑥ LLM Generation (Groq — llama-3.3-70b-versatile) │\n"
"│  ⑦ JSON Extraction & Repair (3-tier strategy)       │\n"
"│  ⑧ Travel Info Injection & Response Assembly        │\n"
"└──────────────────────┬──────────────────────────────┘\n"
"          ┌────────────┴────────────┐\n"
"          ▼                        ▼\n"
"   ChromaDB (local)          Groq Cloud API\n"
"   (destination docs)        (LLM inference)"
)

# ── 5. Technology Stack ───────────────────────────────────────────────────
section_heading(5, "Technology Stack")
sub_heading("5.1 Backend")
make_table(doc,
    ["Component", "Technology", "Purpose"],
    [
        ["Web Framework", "FastAPI + Uvicorn", "REST API, async request handling"],
        ["LLM", "Groq — llama-3.3-70b-versatile", "Itinerary text generation"],
        ["Embeddings", "HuggingFace all-MiniLM-L6-v2", "Semantic text encoding for RAG"],
        ["Vector Database", "ChromaDB (local persistent)", "Destination document retrieval"],
        ["RAG Framework", "LangChain", "Document retrieval and LLM orchestration"],
        ["Flight Data", "Sky-Scrapper RapidAPI + Simulator", "Flight cost estimation"],
        ["Data Processing", "Pandas", "CSV ingestion and document building"],
        ["Validation", "Pydantic", "Request model validation"],
    ]
)

sub_heading("5.2 Frontend")
make_table(doc,
    ["Component", "Technology"],
    [
        ["UI Framework", "React 19 + Vite"],
        ["HTTP Client", "Axios"],
        ["Styling", "Vanilla CSS (inline styles)"],
    ]
)

sub_heading("5.3 Data Sources")
make_table(doc,
    ["CSV File", "Contents"],
    [
        ["Expanded_Destinations.csv", "Indian destinations — name, state, type, popularity, best time"],
        ["India_Tourism_2025_Processed.csv", "State-level tourism revenue, visitor counts, visit purposes"],
        ["Tourist_Destinations.csv", "International destinations — daily cost (USD), season, rating"],
        ["Worldwide_Travel_Cities_Dataset_*.csv", "Global city scores — culture, adventure, nature, cuisine"],
    ]
)

# ── 6. Key Modules ────────────────────────────────────────────────────────
section_heading(6, "Key Modules")

sub_heading("6.1  build_destination_documents()")
body("Reads all four CSV files and produces one rich LangChain Document per destination. "
     "Indian destinations are cross-joined with state tourism statistics. International destinations "
     "include USD/day costs converted to INR at ₹84/USD.")

sub_heading("6.2  is_out_of_domain()")
body("Checks whether retrieved documents actually contain keywords from the queried city name, "
     "preventing irrelevant results from passing source/destination validation.")

sub_heading("6.3  extract_json() + _repair_truncated_json()")
body("Three-tier JSON extraction strategy to handle imperfect LLM outputs:")
bullet("Tier 1 — Parse raw LLM response directly")
bullet("Tier 2 — Repair truncation by tracking open braces/brackets character-by-character and closing them")
bullet("Tier 3 — Truncate to the last complete } as a last resort")

sub_heading("6.4  budget_validate()")
body("Validates the remaining ground budget (after deducting flight cost) against minimum thresholds:")
make_table(doc,
    ["Route Type", "Minimum Budget"],
    [
        ["Domestic India", "₹500 per person per day"],
        ["International", "₹3,000 per person per day"],
    ]
)

sub_heading("6.5  budget_allocate()")
body("Splits the ground budget using standard travel ratios:")
make_table(doc,
    ["Component", "Allocation", "Rationale"],
    [
        ["Accommodation (Stay)", "40%", "Largest single expense for most trips"],
        ["Food & Dining", "25%", "3 meals per day across the trip"],
        ["Activities & Sightseeing", "20%", "Entry fees, guided tours, experiences"],
        ["Local Transport", "15%", "Taxis, buses, rentals within destination"],
    ]
)

sub_heading("6.6  Flight Simulator (flight_data.py)")
body("Detects route type and generates realistic airlines, prices, and durations:")
make_table(doc,
    ["Route Type", "Detection", "Price Range", "Duration"],
    [
        ["Domestic India", "Both cities in Indian cities list", "₹3,500–₹9,500", "1h–3h 20m"],
        ["Neighbour Country", "Sri Lanka, Nepal, Bhutan, Maldives…", "₹12,000–₹35,000", "1h 30m–5h"],
        ["International", "All other destinations", "₹45,000–₹1,20,000", "7h–15h"],
    ]
)
body("Real Sky-Scrapper API (RapidAPI) is attempted first when RAPID_API_KEY is configured; "
     "the simulator is an automatic fallback.")

# ── 7. Request Processing Pipeline ────────────────────────────────────────
section_heading(7, "Request Processing Pipeline")
steps = [
    ("Source Validation", "Source city keywords checked in ChromaDB. Not found → error."),
    ("Flight Cost Lookup", "RapidAPI attempted; simulated fallback applied automatically."),
    ("Budget Guard", "If flight cost ≥ total budget → error with cheapest flight price shown."),
    ("Ground Budget", "Ground Budget = Total Budget − Flight Cost."),
    ("Budget Validation", "Ground budget checked against per-person-per-day minimum."),
    ("Budget Allocation", "Stay 40% / Food 25% / Activities 20% / Transport 15% of ground budget."),
    ("Destination Validation", "Destination keywords checked in ChromaDB. Not found → error."),
    ("RAG Retrieval", "Full destination documents retrieved (k=6). Cost hints extracted per component."),
    ("LLM Generation", "Groq LLM generates JSON itinerary seeded with correct budget values and specificity rules."),
    ("JSON Repair", "3-tier extraction repairs any truncated or malformed output."),
    ("Travel Injection", "Travel info (already computed) injected directly — never passed through LLM."),
]
for i, (step, desc) in enumerate(steps, 1):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_after = Pt(3)
    add_inline_text(p, f"Step {i} — {step}: ", bold=True, size=10.5, color="B87333")
    add_inline_text(p, desc, size=10.5)

# ── 8. API Endpoints ──────────────────────────────────────────────────────
section_heading(8, "API Endpoints")

sub_heading("POST  /init-db")
body("Builds ChromaDB from the static_rag/ CSV files. Must be called once before generating itineraries.")
add_code_block(doc, 'Response: { "status": "success", "documents": 1248 }')

sub_heading("POST  /generate")
body("Accepts trip parameters and returns a complete, structured itinerary.")
add_code_block(doc,
'Request:\n'
'{\n'
'  "source": "Mumbai",\n'
'  "destination": "Goa",\n'
'  "budget": 25000,\n'
'  "days": 3,\n'
'  "food": "non-vegetarian",\n'
'  "travelers": 2\n'
'}\n\n'
'Response:\n'
'{\n'
'  "plan": {\n'
'    "destination": "Goa",\n'
'    "duration_days": 3,\n'
'    "total_budget_inr": 25000,\n'
'    "travelers": 2,\n'
'    "travel": {\n'
'      "from": "Mumbai",\n'
'      "to": "Goa",\n'
'      "options": [\n'
'        { "airline": "IndiGo", "price_per_person_inr": 4200,\n'
'          "total_price_inr": 8400, "duration": "1h 20m" }\n'
'      ]\n'
'    },\n'
'    "trip_plan": [\n'
'      {\n'
'        "day": 1,\n'
'        "theme": "Old Goa Heritage",\n'
'        "activities": [\n'
'          { "time": "Morning", "activity": "Visit Basilica of Bom Jesus", "cost_inr": 0 }\n'
'        ],\n'
'        "food": [\n'
'          { "meal": "Lunch", "description": "Goan Fish Curry at Ritz Classic", "cost_inr": 350 }\n'
'        ],\n'
'        "stay": { "name": "Old Quarter Inn", "type": "Budget", "cost_per_night_inr": 1400 },\n'
'        "transport": { "mode": "Hired scooter rental", "cost_inr": 400 },\n'
'        "day_total_inr": 2920\n'
'      }\n'
'    ],\n'
'    "tips": ["Visit beaches early morning", "Carry cash", "Best season Oct-Feb"]\n'
'  },\n'
'  "error": false\n'
'}'
)

sub_heading("POST  /chat")
body("Accepts follow-up questions with a session_id to maintain conversation context.")
add_code_block(doc,
'Request:  { "message": "Suggest cheaper hotels", "session_id": "abc123" }\n'
'Response: { "response": "For budget stays under Rs.1,500/night...", "error": false }'
)

# ── 9. Input Validations ──────────────────────────────────────────────────
section_heading(9, "Input Validations")
make_table(doc,
    ["Validation Check", "Trigger Condition", "Error Message"],
    [
        ["Source in database", "Source keywords absent from ChromaDB", "Source 'X' not found in our database"],
        ["Destination in database", "Destination keywords absent from ChromaDB", "Destination 'X' not found in our database"],
        ["Flight affordability", "Flight cost >= total budget", "Budget insufficient to cover flight (Rs.Y)"],
        ["Ground budget — domestic", "< Rs.500/person/day after flights", "Budget too low — minimum Rs.500/person/day"],
        ["Ground budget — international", "< Rs.3,000/person/day after flights", "Budget too low — minimum Rs.3,000/person/day"],
        ["Empty LLM response", "Model returns blank content", "AI returned an empty response"],
        ["LLM refusal", "Refusal phrases detected, no { in response", "AI declined this request"],
    ]
)

# ── 10. Challenges & Solutions ────────────────────────────────────────────
section_heading(10, "Challenges & Solutions")
make_table(doc,
    ["Challenge", "Solution Implemented"],
    [
        ["LLM truncating JSON output", "3-tier repair — structural closure, brace tracking, last-} fallback"],
        ["Generic activity/hotel names", "Specificity prompt rules enforce real landmark and hotel names"],
        ["Python dict syntax in JSON", "Travel options injected server-side after parsing — never through LLM"],
        ["Wrong prices for international routes", "Route classifier: domestic / neighbour / international price tiers"],
        ["Budget ignoring flight cost", "Flight computed first; ground budget = total - flight drives all allocations"],
        ["Irrelevant RAG results", "is_out_of_domain() keyword check rejects off-topic documents"],
    ]
)

# ── 11. Frontend Components ───────────────────────────────────────────────
section_heading(11, "Frontend Components (App.jsx)")
make_table(doc,
    ["Component", "Description"],
    [
        ["TravelForm", "Input form — source, destination, budget, days, food preference, travelers"],
        ["ItineraryDisplay", "Renders the full itinerary with summary stat cards"],
        ["TravelCard", "Shows all flight options with Best Price badge on cheapest"],
        ["DayCard", "Expandable day card — activities, meals, stay, transport sections"],
        ["StatCard", "Summary tile showing destination, duration, travelers, total budget"],
        ["Chat Input", "Follow-up question bar with session_id for conversation continuity"],
    ]
)

# ── 12. Setup & Installation ──────────────────────────────────────────────
section_heading(12, "Setup & Installation")

sub_heading("12.1 Prerequisites")
bullet("Python 3.10 or higher")
bullet("Node.js 18+ and npm")
bullet("Groq API key — available at console.groq.com/keys")

sub_heading("12.2 Backend Setup")
add_code_block(doc,
"cd backend\n"
"python -m venv venv\n"
"venv\\Scripts\\activate          # Windows\n"
"# source venv/bin/activate      # macOS / Linux\n\n"
"pip install -r requirements.txt\n\n"
"# Create .env file\n"
"echo GROQ_API_KEY=your_key_here > .env\n\n"
"# Start the server\n"
"uvicorn test2:app --reload\n"
"# Backend available at http://localhost:8000\n\n"
"# Initialize vector database (first run only)\n"
"curl -X POST http://localhost:8000/init-db"
)

sub_heading("12.3 Frontend Setup")
add_code_block(doc,
"cd frontend\n"
"npm install\n"
"npm run dev\n"
"# Frontend available at http://localhost:5173"
)

sub_heading("12.4 Environment Variables")
make_table(doc,
    ["Variable", "Required", "Purpose"],
    [
        ["GROQ_API_KEY", "Yes", "Groq LLM inference — itinerary generation"],
        ["RAPID_API_KEY", "Optional", "Real flight data via Sky-Scrapper API (falls back to simulator)"],
    ]
)

# ── 13. Conclusion ────────────────────────────────────────────────────────
section_heading(13, "Conclusion")
body(
    "This project demonstrates the practical application of Generative AI and Retrieval-Augmented Generation "
    "in the travel domain. By combining structured destination data from curated CSV datasets with the "
    "Groq-hosted llama-3.3-70b-versatile model, the system generates accurate, budget-aware, and specific "
    "travel itineraries."
)
body(
    "The layered validation pipeline — source check, flight cost estimation, ground budget validation, "
    "and RAG retrieval — ensures that only feasible and well-grounded trips are planned. "
    "The 3-tier JSON repair mechanism makes the system robust against LLM output variability. "
    "The clear separation between RAG data (for cost context) and LLM world knowledge (for specific "
    "names and details) produces itineraries that are both data-grounded and richly detailed."
)

# ── 14. Future Enhancements ───────────────────────────────────────────────
section_heading(14, "Future Enhancements")
enhancements = [
    "Real-time hotel booking integration (MakeMyTrip / Booking.com API)",
    "Live weather forecast per destination using a weather API",
    "User accounts with saved and shared itineraries",
    "Multi-city trip routing support",
    "Carbon footprint estimation per transport mode",
    "React Native mobile application",
]
for e in enhancements:
    bullet(e)

# ── Footer ────────────────────────────────────────────────────────────────
doc.add_paragraph()
add_horizontal_rule(doc)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_inline_text(p, "Document prepared for Mini Project evaluation — April 2026", italic=True, size=9.5, color="888888")

# ── Save ──────────────────────────────────────────────────────────────────
output_path = r"d:\Project\Travel Itenary Planner\Mini Project1\genAI\PROJECT_REPORT.docx"
doc.save(output_path)
print(f"Saved: {output_path}")
